import os
import uuid
import io
import re
import shutil
import tempfile
import time
from flask import Flask, request, jsonify, render_template, send_file
try:
    from docx import Document
except ImportError:
    pass # Will be installed
from groq import Groq
from dotenv import load_dotenv

load_dotenv()                             # reads .env
load_dotenv('.env.local', override=True)  # reads .env.local (Vercel CLI generates this)

app = Flask(__name__)

# Max file size: 25MB (Whisper API limit)
app.config['MAX_CONTENT_LENGTH'] = 25 * 1024 * 1024

ALLOWED_EXTENSIONS = {'mp3', 'mp4', 'mpeg', 'mpga', 'm4a', 'wav', 'webm', 'ogg', 'flac'}

# In-memory chunk storage for Vercel (stateless serverless functions
# can't share filesystem between requests)
chunk_storage = {}

# ── Fallback models for chat completions (paraphrase) ──────────────────────────
# Ordered by priority/quality. If primary model hits rate limit, fallback to next.
CHAT_MODELS = [
    "llama-3.3-70b-versatile",        # Primary — best quality
    "llama-3.1-8b-instant",           # Fallback 1 — fast, confirmed active
    "llama-3.1-70b-versatile",        # Fallback 2 — same family, separate quota
    "llama3-8b-8192",                 # Fallback 3 — legacy LLaMA 3, still supported
]


class GroqKeyRotator:
    """
    Manages a pool of Groq API keys.
    When one key hits rate limit (429), automatically rotates to the next.
    """

    def __init__(self):
        self.keys = self._load_keys()
        self.current_index = 0
        if not self.keys:
            raise ValueError(
                "Tidak ada GROQ_API_KEY yang ditemukan. "
                "Tambahkan GROQ_API_KEY (atau GROQ_API_KEY_1, _2, _3, ...) di .env"
            )

    def _load_keys(self) -> list:
        keys = []
        # Support both single key and numbered keys
        single = os.getenv("GROQ_API_KEY", "").strip()
        if single:
            keys.append(single)
        # Numbered keys: GROQ_API_KEY_1, GROQ_API_KEY_2, ...
        for i in range(1, 11):
            k = os.getenv(f"GROQ_API_KEY_{i}", "").strip()
            if k and k not in keys:
                keys.append(k)
        return keys

    def get_client(self) -> Groq:
        return Groq(api_key=self.keys[self.current_index])

    def rotate(self) -> bool:
        """Move to next key. Returns True if rotated, False if all keys exhausted."""
        next_index = (self.current_index + 1) % len(self.keys)
        if next_index == self.current_index or len(self.keys) == 1:
            return False  # Only one key, can't rotate
        self.current_index = next_index
        return True

    def total_keys(self) -> int:
        return len(self.keys)

    def transcribe_with_rotation(self, tmp_path: str, ext: str) -> str:
        """
        Attempt audio transcription, rotating keys on rate limit.
        Returns transcription text or raises Exception.
        """
        tried = set()
        start_index = self.current_index

        while True:
            tried.add(self.current_index)
            try:
                client = self.get_client()
                with open(tmp_path, 'rb') as audio_file:
                    transcript = client.audio.transcriptions.create(
                        model="whisper-large-v3",
                        file=audio_file,
                        language="id",
                        response_format="text"
                    )
                return transcript.strip()
            except Exception as e:
                err = str(e).lower()
                is_rate_limit = 'rate limit' in err or '429' in err or 'rate_limit' in err
                if is_rate_limit and self.rotate():
                    if self.current_index in tried:
                        # All keys tried
                        raise Exception(
                            f'Semua {self.total_keys()} API Key telah mencapai batas. '
                            'Tunggu beberapa menit atau tambahkan API Key baru.'
                        )
                    continue  # Try next key
                raise  # Non-rate-limit error — bubble up

    def paraphrase_with_rotation(self, messages: list, temperature: float, max_tokens: int) -> str:
        """
        Attempt chat completion, rotating keys AND falling back to alternative models.
        Skips models that are rate-limited (429) or decommissioned (400).
        Returns response text or raises Exception.
        """
        tried_combos = set()  # (key_index, model) pairs tried

        for model in CHAT_MODELS:
            # For each model, try all available keys
            key_tried_for_model = set()
            while True:
                combo = (self.current_index, model)
                if combo in tried_combos:
                    break
                tried_combos.add(combo)
                key_tried_for_model.add(self.current_index)

                try:
                    client = self.get_client()
                    completion = client.chat.completions.create(
                        model=model,
                        messages=messages,
                        temperature=temperature,
                        max_tokens=max_tokens,
                    )
                    return completion.choices[0].message.content
                except Exception as e:
                    err = str(e).lower()
                    is_rate_limit = 'rate limit' in err or '429' in err or 'rate_limit' in err
                    is_decommissioned = 'decommissioned' in err or 'model_decommissioned' in err

                    if is_decommissioned:
                        # Model removed by Groq — skip directly to next model
                        break
                    elif is_rate_limit:
                        # Rate limited — try next key for this model
                        rotated = self.rotate()
                        if not rotated or self.current_index in key_tried_for_model:
                            # All keys exhausted for this model → try next model
                            break
                        continue
                    raise  # Other errors (auth, file format, etc.) — bubble up

        # All models + keys exhausted
        raise Exception(
            f'Semua {self.total_keys()} API Key dan {len(CHAT_MODELS)} model cadangan '
            'telah mencapai batas penggunaan. Tunggu beberapa menit atau tambahkan API Key baru.'
        )


# Global key rotator instance
key_rotator = GroqKeyRotator()


def allowed_file(filename: str) -> bool:
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def classify_error(error_msg: str) -> str:
    """Classify Groq API errors into user-friendly messages."""
    msg_lower = error_msg.lower()
    if 'semua' in msg_lower and ('api key' in msg_lower or 'key' in msg_lower):
        return error_msg  # Already user-friendly from rotator
    elif 'api_key' in msg_lower or 'authentication' in msg_lower or 'invalid api key' in msg_lower:
        return 'API Key tidak valid. Periksa GROQ_API_KEY di file .env Anda.'
    elif 'rate limit' in msg_lower or '429' in msg_lower:
        return 'Batas penggunaan API tercapai. Coba lagi beberapa saat.'
    elif 'invalid file format' in msg_lower or 'audio' in msg_lower:
        return 'File audio rusak atau format tidak valid. Pastikan file bisa diputar.'
    elif 'timeout' in msg_lower or 'timed out' in msg_lower:
        return 'Proses terlalu lama. Coba file audio yang lebih pendek.'
    else:
        return f'Terjadi kesalahan: {error_msg}'


@app.route('/')
def index():
    return render_template('index.html')


@app.route('/transcribe', methods=['POST'])
def transcribe():
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'Tidak ada file yang diunggah.'}), 400

        file = request.files['file']

        chunk_index = int(request.form.get('chunkIndex', 0))
        total_chunks = int(request.form.get('totalChunks', 1))
        file_id = request.form.get('fileId', uuid.uuid4().hex)
        original_filename = request.form.get('filename', file.filename)

        if original_filename == '':
            return jsonify({'error': 'Nama file tidak boleh kosong.'}), 400

        if not allowed_file(original_filename):
            ext = original_filename.rsplit('.', 1)[-1].upper() if '.' in original_filename else 'UNKNOWN'
            return jsonify({'error': f'Format file .{ext} tidak didukung.'}), 415

        # Read chunk data into memory (works on Vercel stateless functions)
        chunk_data = file.read()

        if total_chunks == 1:
            # Single chunk — process directly from memory (most common case)
            audio_bytes = chunk_data
        else:
            # Multi-chunk — store in memory dict
            if file_id not in chunk_storage:
                chunk_storage[file_id] = {}
            chunk_storage[file_id][chunk_index] = chunk_data

            # Not all chunks received yet
            if chunk_index < total_chunks - 1:
                return jsonify({'message': f'Chunk {chunk_index} received.'})

            # All chunks received — merge in memory
            audio_bytes = b''
            for i in range(total_chunks):
                if i not in chunk_storage.get(file_id, {}):
                    # Chunk missing — likely hit different Vercel instance
                    # Clean up and return error
                    chunk_storage.pop(file_id, None)
                    return jsonify({
                        'error': 'Upload gagal: beberapa bagian file hilang. '
                                 'Coba upload ulang atau gunakan file yang lebih kecil (maks 4MB untuk deployment online).'
                    }), 400
                audio_bytes += chunk_storage[file_id][i]
            # Clean up stored chunks
            chunk_storage.pop(file_id, None)

        # Write to temp file for Groq API (it needs a file-like object with name)
        ext = original_filename.rsplit('.', 1)[1].lower() if '.' in original_filename else 'mp3'
        tmp_path = None
        try:
            tmp_fd, tmp_path = tempfile.mkstemp(suffix=f'.{ext}')
            with os.fdopen(tmp_fd, 'wb') as tmp_file:
                tmp_file.write(audio_bytes)

            # Use key rotator — automatically switches key on rate limit
            transcript = key_rotator.transcribe_with_rotation(tmp_path, ext)

            return jsonify({'transcription': transcript})

        finally:
            if tmp_path and os.path.exists(tmp_path):
                try:
                    os.remove(tmp_path)
                except:
                    pass

    except Exception as e:
        return jsonify({'error': classify_error(str(e))}), 500


def post_process_naskah(text: str) -> str:
    """
    Ensure consistent spacing in the radio script output.
    Adds blank lines between sections — does NOT modify content or punctuation
    that would break paragraphs.
    """
    if not text:
        return text

    # Remove markdown formatting (**, ##, etc.)
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'^#{1,6}\s+', '', text, flags=re.MULTILINE)

    # Patterns that should have a blank line BEFORE them
    section_patterns = [
        r'\[LEAD.*?\]',
        r'\[DETAIL PENDUKUNG\]',
        r'\[KUTIPAN\]',
        r'\[PENUTUP.*?\]',
        r'Insert\s*-',
        r'\(---\)',
    ]

    # Also ensure blank line before numbered titles like "1// ..." or "2// ..."
    section_patterns.append(r'\d+//')

    for pattern in section_patterns:
        # Add blank line before pattern if not already preceded by one
        # Use negative lookbehind to ensure we don't add duplicate blank lines
        text = re.sub(r'(?<!\n)\n(' + pattern + ')', r'\n\n\1', text)

    # Ensure blank line AFTER titles (titles match \d+// [Title text])
    text = re.sub(r'(\d+//.+?)\n(?!\n)', r'\1\n\n', text)

    # Clean up excessive blank lines (max 1 blank line between sections)
    text = re.sub(r'\n{3,}', '\n\n', text)

    return text.strip()


@app.route('/paraphrase', methods=['POST'])
def paraphrase():
    try:
        data = request.json
        if not data or 'text' not in data:
            return jsonify({'error': 'Teks tidak ditemukan.'}), 400

        text = data['text']

        if not text.strip():
            return jsonify({'error': 'Teks kosong.'}), 400

        messages = [
            {
                "role": "system",
                "content": (
                    "Anda adalah seorang Jurnalis Radio senior yang ahli dalam menulis naskah berita udara (on-air script). "
                    "Tugas Anda adalah mengubah transkrip mentah menjadi BEBERAPA naskah berita radio terpisah yang siap siar "
                    "dengan mengikuti aturan penulisan berikut secara KETAT.\n\n"

                    "## PENANGANAN INPUT TRANSKRIP MENTAH\n"
                    "Input yang Anda terima bisa berupa:\n"
                    "- Transkrip audio dari speech-to-text (tanpa tanda baca / kalimat menyambung / kata berulang)\n"
                    "- Dokumen teks yang sudah rapi\n"
                    "Apapun bentuk input-nya / Anda HARUS:\n"
                    "1. Pahami SELURUH isi transkrip terlebih dahulu //\n"
                    "2. Identifikasi topik-topik utama dan narasumber yang disebutkan //\n"
                    "3. Pisahkan informasi berdasarkan topik / sudut pandang / atau angle berbeda //\n"
                    "4. Tulis ulang menjadi naskah berita yang rapi mengikuti FORMAT di bawah //\n"
                    "5. JANGAN salin kata per kata dari transkrip / tulis ulang dengan bahasa tutur radio //\n"
                    "6. Perbaiki tata bahasa / hilangkan pengulangan / dan susun secara logis //\n\n"

                    "## PRINSIP UTAMA: PIRAMIDA TERBALIK (INVERTED PYRAMID)\n"
                    "Setiap naskah berita WAJIB mengikuti struktur PIRAMIDA TERBALIK:\n"
                    "   PALING PENTING ▼ (Paragraf 1 — Lead)\n"
                    "   PENTING        ▼ (Paragraf 2 — Detail Pendukung)\n"
                    "   CUKUP PENTING  ▼ (Paragraf 3 — Kutipan/Elaborasi)\n"
                    "   TAMBAHAN       ▼ (Paragraf 4 — Latar Belakang/Dampak)\n\n"
                    "Artinya: informasi TERPENTING ditempatkan di awal naskah / dan informasi tambahan di akhir //\n"
                    "Jika pendengar hanya mendengar 10 detik pertama / mereka SUDAH mendapat inti beritanya //\n\n"

                    "## ATURAN UTAMA\n"
                    "1. WAJIB buat MINIMAL 2 berita, idealnya 4 berita dari satu transkrip. "
                    "Pecah berdasarkan topik, sudut pandang, atau aspek berbeda dari isu yang sama.\n"
                    "2. Jika transkrip hanya membahas 1 topik, tetap pecah menjadi minimal 2 berita "
                    "dengan ANGLE (sudut pandang) BERBEDA. Contoh:\n"
                    "   - Berita 1: Fokus pada peristiwa/fakta utama\n"
                    "   - Berita 2: Fokus pada dampak ke masyarakat\n"
                    "   - Berita 3: Fokus pada respons/tanggapan pihak terkait\n"
                    "   - Berita 4: Fokus pada latar belakang/konteks lebih luas\n\n"

                    "## FORMAT JEDA (WAJIB DIPATUHI)\n"
                    "- Gunakan garis miring satu (/) untuk JEDA PENDEK (pengganti koma saat penyiar butuh napas).\n"
                    "- Gunakan garis miring dua (//) untuk AKHIR KALIMAT (pengganti titik).\n"
                    "- Gunakan garis miring tiga (///) sebagai PENUTUP seluruh naskah berita.\n"
                    "- JANGAN PERNAH gunakan tanda titik (.) untuk mengakhiri kalimat / selalu gunakan // sebagai gantinya //\n"
                    "- JANGAN PERNAH gunakan tanda koma (,) jika jeda tersebut untuk napas penyiar / gunakan / sebagai gantinya //\n\n"

                    "## FORMAT JUDUL\n"
                    "Judul menggunakan format: [Nomor Urut]// [Judul Berita]\n"
                    "Judul harus memancing rasa penasaran seperti headline media online:\n"
                    "- Gunakan teknik clickbait yang etis: pertanyaan retoris / fakta mengejutkan / angka spesifik\n"
                    "- Contoh BAGUS: '1// Warga Kaget / Jalan Utama Kota Mendadak Ditutup Tanpa Pemberitahuan'\n"
                    "- Contoh BAGUS: '2// Terungkap / Alasan di Balik Kenaikan Harga Sembako 40 Persen'\n"
                    "- Contoh BURUK: '1// Berita Tentang Jalan Ditutup'\n\n"

                    "## PENYEBUTAN JABATAN\n"
                    "Sebutkan jabatan tokoh SEBELUM namanya.\n"
                    "Contoh: Plt. Bupati Pekalongan / Sukirman //\n"
                    "Contoh: Kepala Dinas Kesehatan / dr. Ahmad //\n\n"

                    "## STRUKTUR NASKAH (FORMAT OUTPUT YANG HARUS DIIKUTI)\n"
                    "Setiap berita WAJIB mengikuti urutan berikut PERSIS seperti ini:\n\n"
                    "```\n"
                    "[Nomor]// [Judul Berita]\n\n"

                    "[LEAD — 5W+1H]\n"
                    "[Paragraf 1 — isi lead / jawab APA SIAPA DI MANA KAPAN //]\n\n"

                    "[DETAIL PENDUKUNG]\n"
                    "[Paragraf 2 — isi detail pendukung / BAGAIMANA / data / kronologi //]\n\n"

                    "[KUTIPAN]\n"
                    "[Paragraf 3 — kutipan tidak langsung dari narasumber //]\n\n"

                    "Insert - [Nama Narasumber] - [Judul Berita Yang Sama Dengan Judul Di Atas]\n\n"

                    "[PENUTUP — TAMBAHAN]\n"
                    "[Paragraf 4 — latar belakang / dampak / rencana //]\n\n"

                    "(---)\n"
                    "///\n"
                    "```\n\n"

                    "## CHECKLIST PIRAMIDA TERBALIK\n"
                    "Sebelum menyelesaikan setiap naskah / pastikan:\n"
                    "- [ ] Paragraf 1 (Lead) sudah menjawab minimal 3 dari 5W+1H //\n"
                    "- [ ] Paragraf 1 bisa BERDIRI SENDIRI sebagai berita utuh //\n"
                    "- [ ] Setiap paragraf berikutnya memiliki tingkat kepentingan LEBIH RENDAH //\n"
                    "- [ ] Jika paragraf 4 dihapus / berita tetap masuk akal //\n"
                    "- [ ] Jika paragraf 3 dan 4 dihapus / pendengar tetap mendapat inti berita //\n\n"

                    "## GAYA BAHASA\n"
                    "- Gunakan bahasa TUTUR yang ringkas / jelas / dan mengalir //\n"
                    "- Hindari kalimat yang terlalu panjang tanpa jeda //\n"
                    "- Tulis seperti penyiar sedang BERBICARA kepada pendengar / bukan membaca laporan //\n"
                    "- Gunakan kalimat aktif / bukan pasif //\n"
                    "- Tetap FAKTUAL — jangan menambah informasi yang tidak ada di transkrip //\n\n"

                    "## CONTOH NASKAH (IKUTI FORMAT INI PERSIS)\n"
                    "```\n"
                    "1// Warga Terkejut / Jalan Protokol Kota Pekalongan Mendadak Ditutup\n\n"

                    "[LEAD — 5W+1H]\n"
                    "Jalan Protokol utama Kota Pekalongan / ditutup total sejak Senin pagi / "
                    "tanpa pemberitahuan resmi kepada warga // Penutupan dipicu oleh proyek "
                    "revitalisasi trotoar yang digagas Pemerintah Kota //\n\n"

                    "[DETAIL PENDUKUNG]\n"
                    "Penutupan diperkirakan berlangsung selama dua pekan / dan berdampak pada "
                    "ribuan pengguna jalan yang melintas setiap harinya // Jalur alternatif "
                    "melalui Jalan Hayam Wuruk dan Jalan Dr. Cipto telah disiapkan / "
                    "namun belum semua warga mengetahuinya //\n\n"

                    "[KUTIPAN]\n"
                    "Kepala Dinas Pekerjaan Umum / Bambang Sutrisno / menyatakan bahwa penutupan ini "
                    "memang harus dilakukan demi kelancaran proyek // Pihaknya mengaku telah mengirimkan "
                    "surat pemberitahuan ke kelurahan setempat / namun informasi tersebut belum sampai "
                    "ke seluruh warga //\n\n"

                    "Insert - Bambang Sutrisno - Warga Terkejut / Jalan Protokol Kota Pekalongan Mendadak Ditutup\n\n"

                    "[PENUTUP — TAMBAHAN]\n"
                    "Proyek revitalisasi ini merupakan bagian dari program nasional penataan "
                    "kota yang sudah direncanakan sejak tahun lalu // Masyarakat berharap agar ke depannya / "
                    "pemerintah dapat memberikan sosialisasi yang lebih luas sebelum melakukan "
                    "penutupan jalan //\n\n"

                    "(---)\n"
                    "///\n"
                    "```\n\n"

                    "## PENTING — WAJIB DIPATUHI\n"
                    "- JANGAN gabung semua topik jadi 1 berita / WAJIB pisahkan //\n"
                    "- Setiap berita harus bisa BERDIRI SENDIRI //\n"
                    "- Setiap berita WAJIB diakhiri dengan (---) lalu /// //\n"
                    "- JANGAN gunakan markdown formatting (** / ## / dll) / tulis plain text saja //\n"
                    "- TIDAK ADA titik (.) dan koma (,) di seluruh naskah / hanya gunakan / dan // //\n\n"

                    "## ATURAN INSERT (WAJIB ADA)\n"
                    "- Setiap berita WAJIB memiliki TEPAT SATU baris Insert //\n"
                    "- Format PASTI: Insert - [Nama Narasumber] - [Judul Berita Ini]\n"
                    "- [Judul Berita Ini] HARUS SAMA PERSIS dengan judul berita yang sedang ditulis //\n"
                    "- Contoh: jika judul berita adalah '1// PMI Gelar Donor Darah Pasca Lebaran' /\n"
                    "  maka tulis: Insert - Mas Fahmi - PMI Gelar Donor Darah Pasca Lebaran //\n"
                    "- Letakkan Insert SETELAH paragraf 3 (kutipan) dan SEBELUM paragraf 4 (penutup) //\n"
                    "- Baris Insert adalah penanda bagi editor untuk menyisipkan audio narasumber //\n"
                    "- Jika tidak ada nama narasumber yang jelas / gunakan jabatan yang disebutkan di transkrip //\n"
                    "- JANGAN PERNAH menghilangkan baris Insert //\n\n"

                    "## ATURAN LABEL SECTION\n"
                    "- Label [LEAD — 5W+1H] / [DETAIL PENDUKUNG] / [KUTIPAN] / [PENUTUP — TAMBAHAN] "
                    "WAJIB ditulis di naskah akhir sebagai penanda section //\n"
                    "- Tulis label persis seperti contoh / dalam tanda kurung siku //\n"
                    "- Label membantu penyiar dan editor memahami struktur naskah //\n\n"

                    "## ENFORCING PIRAMIDA TERBALIK\n"
                    "- Paragraf 1 (Lead) HARUS berisi informasi PALING PENTING — jawab APA / SIAPA / DI MANA / KAPAN //\n"
                    "- Paragraf 2 HARUS berisi detail pendukung — BAGAIMANA / data / kronologi //\n"
                    "- Paragraf 3 HARUS berisi kutipan tidak langsung dari narasumber //\n"
                    "- Paragraf 4 HARUS berisi informasi tambahan — latar belakang / dampak / rencana //\n"
                    "- Tingkat kepentingan HARUS MENURUN dari paragraf 1 ke paragraf 4 //\n"
                    "- Jika paragraf 4 dihapus / berita TETAP masuk akal //\n"
                    "- Jika paragraf 3 dan 4 dihapus / pendengar TETAP mendapat inti berita //"
                )
            },
            {
                "role": "user",
                "content": text
            }
        ]

        # Use key rotator — auto-rotates keys & falls back to alternative models on rate limit
        paraphrased_text = key_rotator.paraphrase_with_rotation(
            messages=messages,
            temperature=0.75,
            max_tokens=5192,
        )
        # Post-process to enforce consistent formatting
        paraphrased_text = post_process_naskah(paraphrased_text)
        return jsonify({'paraphrased': paraphrased_text.strip()})

    except Exception as e:
        return jsonify({'error': classify_error(str(e))}), 500


@app.route('/upload-transcript', methods=['POST'])
def upload_transcript():
    """Accept a .doc/.docx file upload and return the extracted text."""
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'Tidak ada file yang diunggah.'}), 400

        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'Nama file tidak boleh kosong.'}), 400

        filename = file.filename.lower()
        if not (filename.endswith('.docx') or filename.endswith('.doc')):
            return jsonify({'error': 'Format file tidak didukung. Gunakan file .doc atau .docx.'}), 415

        # Read the file content
        file_bytes = file.read()

        try:
            doc = Document(io.BytesIO(file_bytes))
            full_text = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    full_text.append(text)
            extracted = '\n'.join(full_text)
        except Exception:
            return jsonify({'error': 'Gagal membaca file. Pastikan file berformat .docx yang valid.'}), 400

        if not extracted.strip():
            return jsonify({'error': 'File tidak mengandung teks.'}), 400

        return jsonify({'text': extracted.strip(), 'filename': file.filename})

    except Exception as e:
        return jsonify({'error': f'Terjadi kesalahan: {str(e)}'}), 500


@app.route('/export-docx', methods=['POST'])
def export_docx():
    try:
        data = request.json
        if not data or 'text' not in data:
            return jsonify({'error': 'Teks tidak ditemukan.'}), 400

        text = data['text']
        filename = data.get('filename', 'VoiceScript_Document')

        document = Document()
        for paragraph in text.split('\n'):
            if paragraph.strip():
                document.add_paragraph(paragraph.strip())

        f = io.BytesIO()
        document.save(f)
        f.seek(0)

        return send_file(
            f,
            as_attachment=True,
            download_name=f"{filename}.docx",
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        return jsonify({'error': f'Gagal membuat dokumen: {str(e)}'}), 500


@app.errorhandler(413)
def file_too_large(e):
    return jsonify({'error': 'File terlalu besar. Batas maksimal adalah 25MB.'}), 413


if __name__ == '__main__':
    app.run(debug=True, port=5000)
